import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from src.report_generator import ReportBuilder
from src.style_mapper import StyleMapper
from src.template_engine import TemplateEngine


def add_bookmark(paragraph, name: str):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "1")
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "1")
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


class TemplateEngineTests(unittest.TestCase):
    def test_replaces_placeholders_and_removes_template_body_after_anchor(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "template.docx"
            doc = Document()
            doc.add_paragraph("{{MANUAL_NAME}}")
            anchor = doc.add_paragraph("Template body starts here")
            add_bookmark(anchor, "Chapter_0")
            doc.add_paragraph("Old template body")
            doc.save(path)

            out = Path(tmp) / "out.docx"
            engine = TemplateEngine("config/placeholders.json")
            result = engine.load_copy(path, out)
            report = ReportBuilder(str(path), "source.docx", {"manual_name": "Ground Operations Manual"})
            mapper = StyleMapper(result, "config/style_map.json", report)
            engine.apply_metadata(result, {"manual_name": "Ground Operations Manual"}, report, mapper)
            engine.prepare_body_insertion(result, mapper, report)

            texts = [paragraph.text for paragraph in result.paragraphs]
            self.assertIn("Ground Operations Manual", texts)
            self.assertNotIn("Template body starts here", texts)
            self.assertNotIn("Old template body", texts)

    def test_missing_body_insertion_point_stops_conversion(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "template.docx"
            doc = Document()
            doc.add_paragraph("No body bookmark here")
            doc.save(path)

            engine = TemplateEngine("config/placeholders.json")
            result = engine.load_copy(path, Path(tmp) / "out.docx")
            report = ReportBuilder(str(path), "source.docx", {})
            mapper = StyleMapper(result, "config/style_map.json", report)
            with self.assertRaises(ValueError):
                engine.prepare_body_insertion(result, mapper, report)
            self.assertIn("No body content insertion point found in template.", report.to_dict()["errors"])


if __name__ == "__main__":
    unittest.main()
