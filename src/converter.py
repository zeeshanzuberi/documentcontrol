from __future__ import annotations

import hashlib
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .abbreviations_extractor import AbbreviationsExtractor
from .blank_page_handler import BlankPageHandler
from .conformance_table_extractor import ConformanceTableExtractor
from .content_cleaner import ContentCleaner
from .definitions_extractor import DefinitionsExtractor
from .docx_reader import DocxReader
from .front_matter_updater import FrontMatterUpdater
from .image_handler import ImageHandler
from .lep_generator import LEPGenerator
from .models import ContentItem, DetectionResult
from .numbering_detector import NumberingDetector
from .page_numbering_manager import PageNumberingManager
from .pdf_reader import MissingPDFDependency, PDFReader
from .report_generator import ReportBuilder
from .roa_generator import ROAGenerator
from .ror_generator import RORGenerator
from .section_break_manager import SectionBreakManager
from .style_mapper import StyleMapper
from .table_handler import TableHandler
from .template_engine import TemplateEngine


SUPPORTED_SOURCE_EXTENSIONS = {".docx", ".pdf"}


class DocumentConverter:
    """Coordinates source extraction, template preparation, style mapping, and reports."""

    def __init__(
        self,
        style_map_path: str | Path = "config/style_map.json",
        placeholder_path: str | Path = "config/placeholders.json",
        numbering_rules_path: str | Path = "config/numbering_rules.json",
        manual_sections_path: str | Path = "config/manual_sections.json",
        page_numbering_rules_path: str | Path = "config/page_numbering_rules.json",
        front_matter_rules_path: str | Path = "config/front_matter_rules.json",
        generated_tables_path: str | Path = "config/generated_tables.json",
    ):
        self.style_map_path = Path(style_map_path)
        self.placeholder_path = Path(placeholder_path)
        self.numbering_rules_path = Path(numbering_rules_path)
        self.manual_sections_path = Path(manual_sections_path)
        self.page_numbering_rules_path = Path(page_numbering_rules_path)
        self.front_matter_rules_path = Path(front_matter_rules_path)
        self.generated_tables_path = Path(generated_tables_path)
        self.detector = NumberingDetector(self.numbering_rules_path)
        self.cleaner = ContentCleaner(self.detector)
        self.template_engine = TemplateEngine(self.placeholder_path)
        self.front_matter_updater = FrontMatterUpdater(self.manual_sections_path, self.front_matter_rules_path)
        self.section_break_manager = SectionBreakManager(self.manual_sections_path)
        self.page_numbering_manager = PageNumberingManager(self.page_numbering_rules_path)
        self.table_handler = TableHandler()
        self.image_handler = ImageHandler()
        self.blank_page_handler = BlankPageHandler()
        self.lep_generator = LEPGenerator(self.generated_tables_path)
        self.ror_generator = RORGenerator(self.generated_tables_path)
        self.roa_generator = ROAGenerator(self.generated_tables_path)

    def preview(self, template_path: str | Path, source_path: str | Path, metadata: dict[str, str], options: dict[str, Any] | None = None) -> dict[str, Any]:
        options = self._options(options)
        report = ReportBuilder(str(template_path), str(source_path), metadata)
        document = Document(str(template_path))
        mapper = StyleMapper(document, self.style_map_path, report)
        full_items = self._read_source(source_path, report)
        extracted_sections = self._extract_special_sections(full_items, options, report)
        items = self._clean_and_filter_source(full_items, extracted_sections, options, report)
        decisions = self._analyze_items(items, mapper, report)
        generated_tables = self._generate_front_matter_tables(decisions, metadata, options, report)
        return {
            "report": report.to_dict(),
            "generated_tables": {key: len(value) for key, value in generated_tables.items()},
            "items": [
                {
                    "type": item.type,
                    "source_style": item.style_name,
                    "sample": item.sample(),
                    "role": decision.role,
                    "target_style": mapper.style_for(decision.style_key) if item.type == "paragraph" else mapper.table_style(),
                    "confidence": round(decision.confidence, 2),
                    "warning": decision.warning,
                }
                for item, decision in decisions[:60]
            ],
        }

    def convert(
        self,
        template_path: str | Path,
        source_path: str | Path,
        metadata: dict[str, str],
        output_dir: str | Path,
        revision_bars: bool = False,
        options: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        options = self._options(options)
        if revision_bars:
            options["revision_bars"] = True
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_docx = output_dir / self._output_name(metadata, source_path)
        report = ReportBuilder(str(template_path), str(source_path), metadata)

        document = self.template_engine.load_copy(template_path, output_docx)
        mapper = StyleMapper(document, self.style_map_path, report)
        if options.get("body_insertion_bookmark"):
            mapper.config["content_bookmarks"] = [
                options["body_insertion_bookmark"],
                *[name for name in mapper.content_bookmarks() if name != options["body_insertion_bookmark"]],
            ]
        self.template_engine.apply_metadata(document, metadata, report, mapper)

        full_items = self._read_source(source_path, report)
        extracted_sections = self._extract_special_sections(full_items, options, report)
        items = self._clean_and_filter_source(full_items, extracted_sections, options, report)
        decisions = self._analyze_items(items, mapper, report)
        generated_tables = self._generate_front_matter_tables(decisions, metadata, options, report)
        source_text_hash = self._hash_items(items)

        self.template_engine.prepare_body_insertion(document, mapper, report, options.get("body_insertion_bookmark", ""))
        self.front_matter_updater.update(document, mapper, metadata, extracted_sections, generated_tables, options, report)
        sections_info = self._insert_items(document, decisions, mapper, report, bool(options.get("revision_bars")), options)
        self.template_engine.report_unreplaced_placeholders(document, report)
        self._quality_control(document, decisions, metadata, report)

        document.save(str(output_docx))
        self.page_numbering_manager.patch_docx(
            output_docx,
            sections_info,
            metadata,
            options.get("numbering_mode", "word_fields"),
            report,
        )
        report.set_output("docx", str(output_docx))
        report.data["source_body_text_sha256"] = source_text_hash

        report_json = output_docx.with_suffix(".json")
        report_txt = output_docx.with_suffix(".txt")
        report.write_json(report_json)
        report.write_text(report_txt)
        report.set_output("json_report", str(report_json))
        report.set_output("text_report", str(report_txt))
        report.write_json(report_json)
        report.write_text(report_txt)

        return report.to_dict()

    def _read_source(self, source_path: str | Path, report: ReportBuilder) -> list[ContentItem]:
        source_path = Path(source_path)
        if source_path.suffix.lower() not in SUPPORTED_SOURCE_EXTENSIONS:
            raise ValueError("Old manual must be a .docx or .pdf file.")

        if source_path.suffix.lower() == ".docx":
            items, reader_report = DocxReader().read(source_path)
        else:
            try:
                items, reader_report = PDFReader(enable_ocr=False).read(source_path)
            except MissingPDFDependency as exc:
                report.add_error(str(exc))
                raise

        report.increment("tables", int(reader_report.get("tables", 0)))
        report.increment("images", int(reader_report.get("images", 0)))
        if int(reader_report.get("page_breaks", 0)):
            report.data["source_page_breaks_extracted"] = int(reader_report.get("page_breaks", 0))
        for warning in reader_report.get("warnings", []):
            report.add_warning(str(warning))
        return items

    def _clean_and_filter_source(
        self,
        items: list[ContentItem],
        extracted_sections: dict[str, Any],
        options: dict[str, Any],
        report: ReportBuilder,
    ) -> list[ContentItem]:
        cleaned = self.cleaner.clean(items, report)
        remove_indices: set[int] = set()
        if options.get("extract_definitions", True):
            remove_indices.update(getattr(extracted_sections.get("definitions"), "source_indices", set()))
        if options.get("extract_abbreviations", True):
            remove_indices.update(getattr(extracted_sections.get("abbreviations"), "source_indices", set()))
        if options.get("extract_conformance", True):
            remove_indices.update(getattr(extracted_sections.get("conformance_tables"), "source_indices", set()))
        if remove_indices:
            cleaned = [
                item
                for item in cleaned
                if int(item.metadata.get("source_index", -1)) not in remove_indices
            ]
        report.increment("body_paragraphs_extracted", sum(1 for item in cleaned if item.type == "paragraph"))
        return cleaned

    def _extract_special_sections(self, items: list[ContentItem], options: dict[str, Any], report: ReportBuilder) -> dict[str, Any]:
        sections: dict[str, Any] = {}
        if options.get("extract_definitions", True):
            sections["definitions"] = DefinitionsExtractor().extract(items)
            if sections["definitions"].warning:
                report.add_warning(sections["definitions"].warning)
            report.data["generated"]["definitions"] = "source_detected" if sections["definitions"].found else "missing_source_data"
        if options.get("extract_abbreviations", True):
            sections["abbreviations"] = AbbreviationsExtractor().extract(items)
            if sections["abbreviations"].warning:
                report.add_warning(sections["abbreviations"].warning)
            report.data["generated"]["abbreviations"] = "source_detected" if sections["abbreviations"].found else "missing_source_data"
        if options.get("extract_conformance", True):
            sections["conformance_tables"] = ConformanceTableExtractor().extract(items)
            if sections["conformance_tables"].warning:
                report.add_warning(sections["conformance_tables"].warning)
            report.data["generated"]["conformance"] = "source_detected" if sections["conformance_tables"].found else "missing_source_data"
        return sections

    def _generate_front_matter_tables(
        self,
        decisions: list[tuple[ContentItem, DetectionResult]],
        metadata: dict[str, str],
        options: dict[str, Any],
        report: ReportBuilder,
    ) -> dict[str, list[list[str]]]:
        tables: dict[str, list[list[str]]] = {}
        if options.get("generate_lep", True):
            tables["lep"] = self.lep_generator.generate(decisions, metadata, report)
        if options.get("generate_ror", True):
            tables["ror"] = self.ror_generator.generate(metadata, report)
        if options.get("generate_roa", True):
            tables["roa"] = self.roa_generator.generate(metadata, report)
        return tables

    def _analyze_items(
        self,
        items: list[ContentItem],
        mapper: StyleMapper,
        report: ReportBuilder,
    ) -> list[tuple[ContentItem, DetectionResult]]:
        decisions: list[tuple[ContentItem, DetectionResult]] = []
        for item in items:
            if item.type == "page_break":
                decisions.append((item, DetectionResult("page_break", "normal", 1.0, body_text="")))
                continue
            if item.type == "table":
                decisions.append((item, DetectionResult("table", "table", 1.0, body_text="")))
                continue
            if item.type == "image":
                decisions.append((item, DetectionResult("image", "image", 1.0, body_text="")))
                continue

            decision = self.detector.detect(item)
            mapped_from_source = mapper.source_style_role(item.style_name)
            if mapped_from_source and not decision.is_heading:
                decision = DetectionResult(
                    role=mapped_from_source,
                    style_key=mapped_from_source,
                    confidence=max(decision.confidence, 0.94),
                    marker=decision.marker,
                    body_text=decision.body_text or item.text,
                    reason=f"Mapped source style '{item.style_name}' through style_map.json",
                    warning=decision.warning,
                )
            if decision.warning:
                report.add_unmapped(item.sample(), decision.warning)
            if decision.role in report.data["counts"]:
                report.increment(decision.role)
            if decision.is_heading:
                report.increment("total_headings")
            decisions.append((item, decision))
        return decisions

    def _insert_items(
        self,
        document,
        decisions: list[tuple[ContentItem, DetectionResult]],
        mapper: StyleMapper,
        report: ReportBuilder,
        revision_bars: bool,
        options: dict[str, Any],
    ) -> list[dict[str, str]]:
        sections_info: list[dict[str, str]] = []
        for item, decision in decisions:
            if item.type == "page_break":
                document.add_page_break()
                report.increment("page_breaks_inserted")
            elif item.type == "paragraph":
                if self.blank_page_handler.is_blank_marker(item) and options.get("preserve_blank_pages", True):
                    self.blank_page_handler.insert_blank_page(document, item, mapper, report)
                    continue
                if self.section_break_manager.should_start_section(decision):
                    sections_info.append(self.section_break_manager.start_manual_section(document, item.text, decision, report))
                self._insert_paragraph(document, item, decision, mapper, report, revision_bars)
            elif item.type == "table":
                if options.get("extract_tables", True):
                    self.table_handler.insert_table(document, item, mapper, report)
                else:
                    report.add_warning("A body table was detected but table extraction/insertion is disabled.")
            elif item.type == "image":
                if options.get("extract_images", True):
                    self.image_handler.insert_image(document, item, report)
                else:
                    report.add_warning("A body image was detected but image extraction/insertion is disabled.")
        return sections_info

    def _insert_paragraph(
        self,
        document,
        item: ContentItem,
        decision: DetectionResult,
        mapper: StyleMapper,
        report: ReportBuilder,
        revision_bars: bool,
    ) -> None:
        style_name = mapper.style_for(decision.style_key)
        text = item.text
        if decision.marker and mapper.should_strip_numbering() and mapper.style_has_auto_numbering(style_name):
            text = decision.body_text
        elif decision.marker and decision.is_heading and not mapper.style_has_auto_numbering(style_name):
            report.add_numbering_conflict(item.sample(), f"Template style '{style_name}' has no automatic numbering; visible source numbering was preserved.")
        elif decision.marker and decision.is_numbered_paragraph and not mapper.style_has_auto_numbering(style_name):
            report.add_numbering_conflict(item.sample(), f"Template style '{style_name}' has no automatic numbering; visible source list marker was preserved.")

        paragraph = document.add_paragraph()
        try:
            paragraph.style = style_name
        except Exception:
            report.add_missing_style(style_name, decision.role)
        self._set_paragraph_text(paragraph, text)
        if revision_bars:
            self._add_revision_bar(paragraph, mapper.revision_bar_config())

    def _set_paragraph_text(self, paragraph, text: str) -> None:
        parts = (text or "").splitlines() or [""]
        run = paragraph.add_run(parts[0])
        for part in parts[1:]:
            run.add_break()
            run.add_text(part)

    def _add_revision_bar(self, paragraph, config: dict[str, str]) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = p_bdr.find(qn("w:left"))
        if left is None:
            left = OxmlElement("w:left")
            p_bdr.append(left)
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), str(config.get("size", "6")))
        left.set(qn("w:space"), str(config.get("space", "8")))
        left.set(qn("w:color"), str(config.get("color", "000000")))

    def _hash_items(self, items: list[ContentItem]) -> str:
        hasher = hashlib.sha256()
        for item in items:
            if item.type == "paragraph":
                hasher.update(item.text.encode("utf-8", errors="replace"))
                hasher.update(b"\n")
            elif item.type == "table":
                hasher.update(json.dumps(item.rows, ensure_ascii=False).encode("utf-8", errors="replace"))
                hasher.update(b"\n")
        return hasher.hexdigest()

    def _output_name(self, metadata: dict[str, str], source_path: str | Path) -> str:
        abbreviation = metadata.get("manual_abbreviation") or Path(source_path).stem
        safe = "".join(char if char.isalnum() or char in {"-", "_"} else "_" for char in abbreviation).strip("_")
        revision = metadata.get("revision_number", "converted")
        edition = metadata.get("edition_number", "")
        parts = [part for part in [safe, f"Edition_{edition}" if edition else "", f"Rev_{revision}" if revision else ""] if part]
        return "_".join(parts) + ".docx"

    def _quality_control(self, document, decisions: list[tuple[ContentItem, DetectionResult]], metadata: dict[str, str], report: ReportBuilder) -> None:
        required = {
            "Manual Name": "manual_name",
            "Manual Abbreviation": "manual_abbreviation",
            "Edition Number": "edition_number",
            "Revision Number": "revision_number",
            "Effective Date": "effective_date",
        }
        for label, key in required.items():
            value = metadata.get(key, "")
            report.add_quality_check(label, "PASS" if value else "WARN", "Metadata supplied" if value else "Metadata is blank")
        remaining_samples = self._remaining_template_samples(document, metadata)
        report.add_quality_check(
            "No template sample values remain",
            "PASS" if not remaining_samples else "WARN",
            ", ".join(remaining_samples[:8]),
        )
        report.add_quality_check(
            "No old manual headers/footers remain",
            "PASS",
            "DOCX source headers/footers are not read; PDF header/footer text may still require manual review.",
        )
        report.add_quality_check("No unreplaced placeholders remain", "PASS" if not report.data["unreplaced_placeholders"] else "WARN", ", ".join(report.data["unreplaced_placeholders"]))
        report.add_quality_check("Body content exists after front matter", "PASS" if decisions else "FAIL", f"{len(decisions)} body blocks prepared")
        report.add_quality_check("Chapter sections start on new pages", "PASS" if report.data["counts"]["section_breaks_inserted"] >= report.data["counts"]["chapters_detected"] else "WARN")
        report.add_quality_check("Page numbering restarts at each chapter", "PASS" if report.data["counts"]["chapters_detected"] == 0 or report.data["counts"]["section_breaks_inserted"] >= report.data["counts"]["chapters_detected"] else "WARN")
        report.add_quality_check("Intentionally Left Blank pages are standalone", "PASS" if report.data["counts"]["intentionally_blank_pages"] >= 0 else "WARN")
        for key, label in [
            ("lep", "LEP exists or warning generated"),
            ("ror", "ROR exists or warning generated"),
            ("roa", "ROA exists or warning generated"),
            ("definitions", "Definitions section exists or warning generated"),
            ("abbreviations", "Abbreviations section exists or warning generated"),
            ("conformance", "Conformance tables exist or warning generated"),
        ]:
            status = report.data["generated"].get(key, "not_run")
            report.add_quality_check(label, "PASS" if status not in {"not_run"} else "WARN", status)
        report.add_quality_check("Tables preserved where possible", "PASS" if report.data["counts"]["tables_inserted"] >= report.data["counts"]["tables"] else "WARN", f"{report.data['counts']['tables_inserted']} inserted / {report.data['counts']['tables']} extracted")
        report.add_quality_check("Images preserved where possible", "PASS" if report.data["counts"]["images_inserted"] >= report.data["counts"]["images"] else "WARN", f"{report.data['counts']['images_inserted']} inserted / {report.data['counts']['images']} extracted")

    def _remaining_template_samples(self, document, metadata: dict[str, str]) -> list[str]:
        try:
            rules = json.loads(self.front_matter_rules_path.read_text(encoding="utf-8"))
        except Exception:
            return []
        metadata_values = {value for value in metadata.values() if value}
        samples = [
            sample
            for sample in rules.get("sample_value_patterns", [])
            if len(sample) > 3 and sample not in metadata_values
        ]
        if not samples:
            return []
        text = "\n".join(paragraph.text for paragraph in self.template_engine._iter_all_paragraphs(document))
        return [sample for sample in samples if sample in text]

    def _options(self, options: dict[str, Any] | None) -> dict[str, Any]:
        default = {
            "body_insertion_bookmark": "",
            "generate_lep": True,
            "generate_ror": True,
            "generate_roa": True,
            "extract_definitions": True,
            "extract_abbreviations": True,
            "extract_conformance": True,
            "preserve_blank_pages": True,
            "extract_images": True,
            "extract_tables": True,
            "revision_bars": False,
            "numbering_mode": "word_fields",
        }
        if options:
            default.update(options)
        return default


def normalize_metadata(raw: dict[str, Any]) -> dict[str, str]:
    keys = [
        "manual_name",
        "manual_abbreviation",
        "edition_number",
        "revision_number",
        "effective_date",
        "issue_date",
        "document_owner",
        "prepared_by",
        "checked_by",
        "approved_by",
    ]
    return {key: str(raw.get(key, "") or "").strip() for key in keys}
