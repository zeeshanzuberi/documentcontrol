from __future__ import annotations

from io import BytesIO

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu, Inches

from .models import ContentItem
from .report_generator import ReportBuilder


class ImageHandler:
    """Places source body images near their original location."""

    def __init__(self, max_width_inches: float = 6.5):
        self.max_width = Inches(max_width_inches)

    def insert_image(self, document, item: ContentItem, report: ReportBuilder) -> None:
        if not item.image_bytes:
            return
        try:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            width = self._safe_width(item)
            if width:
                run.add_picture(BytesIO(item.image_bytes), width=width)
            else:
                run.add_picture(BytesIO(item.image_bytes), width=self.max_width)
            report.increment("images_inserted")
        except Exception as exc:
            report.add_warning(f"An image could not be inserted and needs manual review: {exc}")

    def _safe_width(self, item: ContentItem):
        width_emu = item.metadata.get("width_emu")
        if not width_emu:
            return self.max_width
        width = Emu(int(width_emu))
        return width if width <= self.max_width else self.max_width
