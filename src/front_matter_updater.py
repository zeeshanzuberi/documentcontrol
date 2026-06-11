from __future__ import annotations

import json
import re
from pathlib import Path

from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from .models import ContentItem, ExtractedSection
from .report_generator import ReportBuilder
from .style_mapper import StyleMapper
from .table_handler import TableHandler


class FrontMatterUpdater:
    """Replaces template example front matter with generated/manual-specific content."""

    def __init__(
        self,
        manual_sections_path: str | Path = "config/manual_sections.json",
        front_matter_rules_path: str | Path = "config/front_matter_rules.json",
    ):
        self.manual_sections = json.loads(Path(manual_sections_path).read_text(encoding="utf-8"))
        self.rules = json.loads(Path(front_matter_rules_path).read_text(encoding="utf-8"))
        self.table_handler = TableHandler()

    def update(
        self,
        document,
        mapper: StyleMapper,
        metadata: dict[str, str],
        extracted_sections: dict[str, ExtractedSection],
        generated_tables: dict[str, list[list[str]]],
        options: dict,
        report: ReportBuilder,
    ) -> None:
        self._update_approval_page(document, mapper, metadata, report)
        if options.get("generate_lep", True):
            self._replace_with_table(document, "lep", generated_tables.get("lep", []), mapper, report)
        if options.get("generate_ror", True):
            self._replace_with_table(document, "ror", generated_tables.get("ror", []), mapper, report)
        if options.get("generate_roa", True):
            self._replace_with_table(document, "roa", generated_tables.get("roa", []), mapper, report)
        if options.get("extract_definitions", True):
            self._replace_with_extracted_section(document, "definitions", extracted_sections.get("definitions"), mapper, report)
        if options.get("extract_abbreviations", True):
            self._replace_with_extracted_section(document, "abbreviations", extracted_sections.get("abbreviations"), mapper, report)
        if options.get("extract_conformance", True):
            self._replace_with_extracted_section(document, "conformance", extracted_sections.get("conformance_tables"), mapper, report)

    def _update_approval_page(self, document, mapper: StyleMapper, metadata: dict[str, str], report: ReportBuilder) -> None:
        approval = self._find_heading(document, self._heading_patterns("approval"))
        if approval is None:
            report.add_warning("Approval Page heading was not found in the template.")
            return
        table = self._first_table_after(document, approval)
        if table is not None:
            values = {
                "prepared": metadata.get("prepared_by", ""),
                "checked": metadata.get("checked_by", ""),
                "approved": metadata.get("approved_by", ""),
                "manual": metadata.get("manual_name", ""),
                "department": metadata.get("document_owner", ""),
            }
            for row in table.rows:
                label = " ".join(row.cells[0].text.lower().split()) if row.cells else ""
                if "prepared" in label and len(row.cells) > 1:
                    row.cells[1].text = values["prepared"]
                elif ("checked" in label or "review" in label) and len(row.cells) > 1:
                    row.cells[1].text = values["checked"]
                elif "approved" in label and len(row.cells) > 1:
                    row.cells[1].text = values["approved"]
                if len(row.cells) > 2 and values["department"]:
                    if "designation" in " ".join(cell.text.lower() for cell in row.cells[:1]) or row.cells[2].text.strip():
                        row.cells[2].text = values["department"]
        rows = [["Field", "Value"]]
        for label, field in self.rules.get("approval", {}).get("fields", []):
            rows.append([label, metadata.get(field, "")])
        self._insert_table_after(document, approval, rows, mapper, report)
        report.data["generated"]["approval"] = "updated_from_metadata"

    def _replace_with_table(self, document, section_key: str, rows: list[list[str]], mapper: StyleMapper, report: ReportBuilder) -> None:
        if not rows:
            report.add_warning(f"{section_key.upper()} generation was enabled but no rows were produced.")
            return
        heading = self._find_or_create_heading(document, section_key, mapper)
        self._clear_region_after_heading(document, heading)
        self._insert_table_after(document, heading, rows, mapper, report)

    def _replace_with_extracted_section(
        self,
        document,
        section_key: str,
        extracted: ExtractedSection | None,
        mapper: StyleMapper,
        report: ReportBuilder,
    ) -> None:
        heading = self._find_or_create_heading(document, section_key, mapper)
        self._clear_region_after_heading(document, heading)
        if not extracted or not extracted.found or not extracted.items:
            status = extracted.warning if extracted else f"No {section_key} section detected in source manual."
            report.add_warning(status)
            report.data["generated"][section_key] = "missing_source_data"
            return
        cursor = heading._p
        for item in extracted.items:
            if item.type == "paragraph" and item.text.strip():
                paragraph = self._new_paragraph(document, item.text, mapper.style_for("definition" if section_key == "definitions" else "body"))
                cursor.addnext(paragraph._p)
                cursor = paragraph._p
            elif item.type == "table" and item.rows:
                table = self._new_table(document, item, mapper, report)
                cursor.addnext(table._tbl)
                cursor = table._tbl
                report.increment("tables_inserted")
        report.data["generated"][section_key] = "transferred_from_source"

    def _find_or_create_heading(self, document, section_key: str, mapper: StyleMapper) -> Paragraph:
        heading = self._find_heading(document, self._heading_patterns(section_key))
        if heading is not None:
            return heading
        style = mapper.style_for("section_heading")
        label = self._label_for(section_key)
        paragraph = document.add_paragraph(label)
        try:
            paragraph.style = style
        except Exception:
            pass
        return paragraph

    def _clear_region_after_heading(self, document, heading: Paragraph) -> None:
        body = document.element.body
        children = list(body)
        try:
            start = children.index(heading._p) + 1
        except ValueError:
            return
        for child in children[start:]:
            if child.tag == qn("w:sectPr"):
                break
            if child.tag == qn("w:p"):
                paragraph = Paragraph(child, document)
                if self._is_front_matter_boundary(paragraph):
                    break
                if child.xpath(".//w:sectPr"):
                    continue
            elif child.tag != qn("w:tbl"):
                continue
            body.remove(child)

    def _insert_table_after(self, document, heading: Paragraph, rows: list[list[str]], mapper: StyleMapper, report: ReportBuilder) -> None:
        table = self.table_handler.create_generated_table(document, rows, mapper, report)
        body = document.element.body
        body.remove(table._tbl)
        heading._p.addnext(table._tbl)
        report.increment("tables_inserted")

    def _new_table(self, document, item: ContentItem, mapper: StyleMapper, report: ReportBuilder) -> Table:
        if item.metadata.get("table_xml"):
            tbl = parse_xml(str(item.metadata["table_xml"]))
            body = document.element.body
            body.insert(len(body) - 1, tbl)
            table = Table(tbl, document)
            try:
                table.style = mapper.table_style()
            except Exception:
                pass
            body.remove(tbl)
            return table
        table = self.table_handler.create_generated_table(document, item.rows, mapper, report)
        document.element.body.remove(table._tbl)
        return table

    def _new_paragraph(self, document, text: str, style_name: str) -> Paragraph:
        paragraph = document.add_paragraph(text)
        try:
            paragraph.style = style_name
        except Exception:
            pass
        body = document.element.body
        body.remove(paragraph._p)
        return paragraph

    def _find_heading(self, document, patterns: list[str]) -> Paragraph | None:
        preferred: list[Paragraph] = []
        fallback: list[Paragraph] = []
        for paragraph in document.paragraphs:
            text = " ".join((paragraph.text or "").split())
            if not text:
                continue
            if any(re.search(pattern, text, flags=re.I) for pattern in patterns):
                if paragraph.style and paragraph.style.name in {"Center Heading 1", "Chapter Heading", "Heading 1"}:
                    preferred.append(paragraph)
                else:
                    fallback.append(paragraph)
        return (preferred or fallback or [None])[0]

    def _first_table_after(self, document, heading: Paragraph) -> Table | None:
        body = document.element.body
        children = list(body)
        try:
            start = children.index(heading._p) + 1
        except ValueError:
            return None
        for child in children[start:]:
            if child.tag == qn("w:p"):
                paragraph = Paragraph(child, document)
                if self._is_front_matter_boundary(paragraph):
                    return None
            if child.tag == qn("w:tbl"):
                return Table(child, document)
        return None

    def _is_front_matter_boundary(self, paragraph: Paragraph) -> bool:
        text = " ".join((paragraph.text or "").split())
        if not text:
            return False
        if paragraph.style and paragraph.style.name in {"Center Heading 1", "Chapter Heading"}:
            return True
        all_patterns = []
        for patterns in self.manual_sections.get("front_matter_headings", {}).values():
            all_patterns.extend(patterns)
        normalized = text.strip().lower()
        return any(
            normalized == pattern.lower()
            or normalized == f"{pattern.lower()}s"
            or normalized.startswith(f"{pattern.lower()} (")
            for pattern in all_patterns
        )

    def _heading_patterns(self, section_key: str) -> list[str]:
        patterns = self.manual_sections.get("front_matter_headings", {}).get(section_key, [])
        return [re.escape(pattern) for pattern in patterns] or [re.escape(self._label_for(section_key))]

    def _label_for(self, section_key: str) -> str:
        return {
            "lep": "List of Effective Pages (LEP)",
            "ror": "Record of Revisions (ROR)",
            "roa": "Record of Amendments (ROA)",
            "definitions": "Definitions",
            "abbreviations": "Abbreviations",
            "conformance": "Conformance Tables",
            "approval": "Approval Page",
        }.get(section_key, section_key.replace("_", " ").title())
