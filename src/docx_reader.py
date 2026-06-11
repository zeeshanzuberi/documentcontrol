from __future__ import annotations

from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from .models import ContentItem


def iter_block_items(parent: DocxDocument | _Cell) -> Iterator[Paragraph | Table]:
    """Yield paragraphs and tables in document order from a document or cell."""

    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


class DocxReader:
    """Extracts body paragraphs, tables, and inline images from old DOCX files."""

    def read(self, path: str | Path) -> tuple[list[ContentItem], dict[str, int | list[str]]]:
        document = Document(str(path))
        items: list[ContentItem] = []
        image_count = 0
        table_count = 0
        page_breaks = 0

        for block in iter_block_items(document):
            if isinstance(block, Paragraph):
                text = block.text or ""
                if self._has_page_break_before(block):
                    page_breaks += 1
                    items.append(ContentItem(type="page_break", metadata={"source": "pageBreakBefore"}))
                if text.strip():
                    items.append(
                        ContentItem(
                            type="paragraph",
                            text=text,
                            style_name=block.style.name if block.style else "",
                            metadata=self._paragraph_metadata(block),
                        )
                    )
                if self._has_explicit_page_break(block):
                    page_breaks += 1
                    items.append(ContentItem(type="page_break", metadata={"source": "run_page_break"}))
                for image in self._images_from_paragraph(block):
                    image_count += 1
                    items.append(image)
            else:
                rows = self._table_rows(block)
                if rows:
                    table_count += 1
                    table_images = self._count_table_images(block)
                    image_count += table_images
                    items.append(
                        ContentItem(
                            type="table",
                            rows=rows,
                            metadata={
                                "source_style": block.style.name if block.style else "",
                                "table_xml": block._tbl.xml,
                                "has_merged_cells": self._has_merged_cells(block),
                                "images_in_table": table_images,
                                "column_count": self._grid_column_count(block),
                            },
                        )
                    )

        for index, item in enumerate(items):
            item.metadata["source_index"] = index

        return items, {"tables": table_count, "images": image_count, "page_breaks": page_breaks, "warnings": []}

    def _table_rows(self, table: Table) -> list[list[str]]:
        rows: list[list[str]] = []
        for row in table.rows:
            row_values: list[str] = []
            for cell in row.cells:
                paragraphs = [paragraph.text for paragraph in cell.paragraphs]
                row_values.append("\n".join(paragraphs).strip())
            if any(value for value in row_values):
                rows.append(row_values)
        return rows

    def _paragraph_metadata(self, paragraph: Paragraph) -> dict[str, object]:
        text = paragraph.text or ""
        runs = [run for run in paragraph.runs if run.text]
        bold_runs = [run for run in runs if run.bold]
        letters = [char for char in text if char.isalpha()]
        all_caps = bool(letters) and sum(1 for char in letters if char.isupper()) / len(letters) > 0.7

        num_pr = paragraph._p.xpath("./w:pPr/w:numPr")
        return {
            "bold": bool(runs and len(bold_runs) / len(runs) >= 0.5),
            "all_caps": all_caps,
            "has_word_numbering": bool(num_pr),
            "spacing_before": paragraph.paragraph_format.space_before.pt if paragraph.paragraph_format.space_before else None,
            "spacing_after": paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else None,
            "alignment": str(paragraph.alignment) if paragraph.alignment is not None else "",
        }

    def _images_from_paragraph(self, paragraph: Paragraph) -> list[ContentItem]:
        images: list[ContentItem] = []
        blips = paragraph._p.xpath(".//a:blip")
        for blip in blips:
            rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if not rel_id:
                continue
            part = paragraph.part.related_parts.get(rel_id)
            if not part:
                continue
            content_type = getattr(part, "content_type", "image/png")
            ext = "." + content_type.split("/")[-1].replace("jpeg", "jpg")
            width_emu, height_emu = self._image_extent(blip)
            images.append(
                ContentItem(
                    type="image",
                    image_bytes=part.blob,
                    image_ext=ext,
                    metadata={
                        "content_type": content_type,
                        "width_emu": width_emu,
                        "height_emu": height_emu,
                    },
                )
            )
        return images

    def _count_table_images(self, table: Table) -> int:
        count = 0
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count += len(paragraph._p.xpath(".//a:blip"))
        return count

    def _has_page_break_before(self, paragraph: Paragraph) -> bool:
        return bool(paragraph._p.xpath("./w:pPr/w:pageBreakBefore"))

    def _has_explicit_page_break(self, paragraph: Paragraph) -> bool:
        return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))

    def _has_merged_cells(self, table: Table) -> bool:
        return bool(table._tbl.xpath(".//w:gridSpan | .//w:vMerge"))

    def _grid_column_count(self, table: Table) -> int:
        grid_cols = table._tbl.xpath("./w:tblGrid/w:gridCol")
        if grid_cols:
            return len(grid_cols)
        return max((len(row.cells) for row in table.rows), default=0)

    def _image_extent(self, blip) -> tuple[int | None, int | None]:
        inline_extents = blip.xpath("ancestor::wp:inline[1]/wp:extent")
        anchor_extents = blip.xpath("ancestor::wp:anchor[1]/wp:extent")
        extents = inline_extents or anchor_extents
        if not extents:
            return None, None
        cx = extents[0].get("cx")
        cy = extents[0].get("cy")
        return int(cx) if cx and cx.isdigit() else None, int(cy) if cy and cy.isdigit() else None
