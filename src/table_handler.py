from __future__ import annotations

from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.table import Table

from .models import ContentItem
from .report_generator import ReportBuilder
from .style_mapper import StyleMapper


class TableHandler:
    """Preserves and restyles tables without converting them into plain text."""

    def insert_table(self, document, item: ContentItem, mapper: StyleMapper, report: ReportBuilder) -> None:
        if item.metadata.get("table_xml"):
            self._insert_xml_table(document, item, mapper, report)
        else:
            self._insert_rebuilt_table(document, item, mapper, report)
        report.increment("tables_inserted")

    def create_generated_table(self, document, rows: list[list[str]], mapper: StyleMapper, report: ReportBuilder):
        max_cols = max((len(row) for row in rows), default=1)
        table = document.add_table(rows=0, cols=max_cols)
        self._apply_table_style(table, mapper, report)
        self._fill_table(table, rows, mapper, report)
        return table

    def _insert_xml_table(self, document, item: ContentItem, mapper: StyleMapper, report: ReportBuilder) -> None:
        table_xml = str(item.metadata["table_xml"])
        tbl = parse_xml(table_xml)
        if item.metadata.get("images_in_table"):
            self._strip_table_drawings(tbl)
            report.add_warning(
                "A source table contained embedded images. Table text and merged cells were preserved, "
                "but images inside table cells need manual review because Word relationships cannot be safely copied."
            )
        body = document.element.body
        body.insert(len(body) - 1, tbl)
        table = Table(tbl, document)
        self._apply_table_style(table, mapper, report)
        document.add_paragraph()

    def _insert_rebuilt_table(self, document, item: ContentItem, mapper: StyleMapper, report: ReportBuilder) -> None:
        if not item.rows:
            return
        table = self.create_generated_table(document, item.rows, mapper, report)
        self._apply_table_style(table, mapper, report)
        document.add_paragraph()

    def _fill_table(self, table, rows: list[list[str]], mapper: StyleMapper, report: ReportBuilder) -> None:
        max_cols = max((len(row) for row in rows), default=1)
        table_text_style = mapper.table_text_style()
        for row_values in rows:
            row = table.add_row()
            for index in range(max_cols):
                value = row_values[index] if index < len(row_values) else ""
                cell = row.cells[index]
                cell.text = value
                if table_text_style:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.style = table_text_style
                        except Exception:
                            report.add_missing_style(table_text_style, "table_text")

    def _apply_table_style(self, table, mapper: StyleMapper, report: ReportBuilder) -> None:
        style = mapper.table_style()
        if not style:
            return
        try:
            table.style = style
        except Exception:
            report.add_missing_style(style, "table")

    def _strip_table_drawings(self, tbl) -> None:
        for drawing in list(tbl.xpath(".//w:drawing")):
            parent = drawing.getparent()
            if parent is not None:
                parent.remove(drawing)


def remove_table(table) -> None:
    parent = table._tbl.getparent()
    if parent is not None:
        parent.remove(table._tbl)
