from __future__ import annotations

import json
import re
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from .report_generator import ReportBuilder
from .style_mapper import StyleMapper


class TemplateEngine:
    """Applies metadata and prepares the uploaded template for body insertion."""

    def __init__(self, placeholders_path: str | Path = "config/placeholders.json"):
        self.placeholders_path = Path(placeholders_path)
        self.config = json.loads(self.placeholders_path.read_text(encoding="utf-8"))

    def load_copy(self, template_path: str | Path, output_path: str | Path) -> DocxDocument:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_bytes(Path(template_path).read_bytes())
        return Document(str(output_path))

    def apply_metadata(self, document: DocxDocument, metadata: dict[str, str], report: ReportBuilder, mapper: StyleMapper) -> None:
        self._replace_placeholders(document, metadata, report)
        self._replace_bookmarks(document, metadata, report)
        self._replace_front_matter_patterns(document, metadata, report, mapper.content_bookmarks())
        self._update_document_properties(document, metadata)
        if mapper.update_fields_on_open():
            self._set_update_fields_on_open(document)

    def prepare_body_insertion(self, document: DocxDocument, mapper: StyleMapper, report: ReportBuilder, bookmark_override: str = "") -> None:
        body = document.element.body
        bookmark_names = mapper.content_bookmarks()
        if bookmark_override:
            bookmark_names = [bookmark_override, *[name for name in bookmark_names if name != bookmark_override]]
        anchor = self._find_content_anchor(document, bookmark_names, mapper.content_placeholders())
        if anchor is None:
            report.add_error("No body content insertion point found in template.")
            raise ValueError("No body content insertion point found in template.")

        children = list(body)
        try:
            anchor_index = children.index(anchor)
        except ValueError:
            report.add_error("No body content insertion point found in template.")
            raise ValueError("No body content insertion point found in template.")

        if mapper.replace_existing_body():
            for child in children[anchor_index:]:
                if child.tag == qn("w:sectPr"):
                    continue
                body.remove(child)
        else:
            # Keep the existing body and insert at the end. This avoids disrupting unknown section breaks.
            report.add_warning("Configured to keep existing template body; converted content was appended after existing content.")

    def report_unreplaced_placeholders(self, document: DocxDocument, report: ReportBuilder) -> None:
        pattern = re.compile(r"\{\{[^}]+\}\}")
        for paragraph in self._iter_all_paragraphs(document):
            for placeholder in pattern.findall(paragraph.text or ""):
                report.add_unreplaced_placeholder(placeholder)

    def _replace_placeholders(self, document: DocxDocument, metadata: dict[str, str], report: ReportBuilder) -> None:
        replacements = {
            placeholder: metadata.get(field, "")
            for placeholder, field in self.config.get("placeholders", {}).items()
            if metadata.get(field, "") != ""
        }
        found = {placeholder: False for placeholder in self.config.get("placeholders", {})}
        for paragraph in self._iter_all_paragraphs(document):
            text = paragraph.text
            for placeholder in found:
                if placeholder in text:
                    found[placeholder] = True
            self._replace_in_paragraph(paragraph, replacements)

        for placeholder, was_found in found.items():
            field = self.config.get("placeholders", {}).get(placeholder, "")
            if metadata.get(field) and not was_found:
                report.add_missing_placeholder(placeholder)

    def _replace_bookmarks(self, document: DocxDocument, metadata: dict[str, str], report: ReportBuilder) -> None:
        configured = self.config.get("bookmarks", {})
        found = {bookmark: False for bookmark in configured}
        for paragraph in self._iter_all_paragraphs(document):
            for start in paragraph._p.xpath(".//w:bookmarkStart"):
                name = start.get(qn("w:name"))
                if name not in configured:
                    continue
                found[name] = True
                value = metadata.get(configured[name], "")
                if value:
                    if not self._set_bookmark_text_same_paragraph(paragraph, start, value):
                        report.add_warning(f"Bookmark '{name}' could not be updated safely because it spans unsupported XML.")
        for bookmark, was_found in found.items():
            field = configured[bookmark]
            if metadata.get(field) and not was_found:
                report.add_missing_bookmark(bookmark)

    def _replace_front_matter_patterns(self, document: DocxDocument, metadata: dict[str, str], report: ReportBuilder, content_bookmarks: list[str]) -> None:
        # This runs before source body insertion, so broad template replacement cannot rewrite source content.
        paragraphs = list(self._iter_all_paragraphs(document))
        for rule in self.config.get("front_matter_patterns", []):
            for paragraph in paragraphs:
                if "literal" in rule:
                    field = rule.get("field", "")
                    value = metadata.get(field, "")
                    if not value:
                        continue
                    replacement = rule.get("format", "{value}").format(value=value, **metadata)
                    self._replace_in_paragraph(paragraph, {rule["literal"]: replacement})
                elif "regex" in rule:
                    replacement = self._safe_format(rule.get("template", ""), metadata)
                    if replacement.strip():
                        self._replace_regex_in_paragraph(paragraph, rule["regex"], replacement)

    def _update_document_properties(self, document: DocxDocument, metadata: dict[str, str]) -> None:
        props = self.config.get("document_properties", {})
        core = document.core_properties
        if props.get("title"):
            core.title = self._format_property(props["title"], metadata)
        if props.get("subject"):
            core.subject = self._format_property(props["subject"], metadata)
        if props.get("author"):
            core.author = self._format_property(props["author"], metadata)
        if props.get("keywords"):
            core.keywords = self._format_property(props["keywords"], metadata)
        if props.get("last_modified_by"):
            core.last_modified_by = self._format_property(props["last_modified_by"], metadata)
        if props.get("company"):
            # Company is an extended property. python-docx does not expose a stable setter.
            pass

    def _set_update_fields_on_open(self, document: DocxDocument) -> None:
        settings = document.settings._element
        existing = settings.xpath("./w:updateFields")
        if existing:
            existing[0].set(qn("w:val"), "true")
            return
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)

    def _find_content_anchor(self, document: DocxDocument, bookmark_names: list[str], placeholders: list[str]):
        body = document.element.body
        bookmark_set = set(bookmark_names)
        for child in body.iterchildren():
            for start in child.xpath(".//w:bookmarkStart"):
                if start.get(qn("w:name")) in bookmark_set:
                    return child

        for child in body.iterchildren():
            if child.tag != qn("w:p"):
                continue
            paragraph = Paragraph(child, document)
            if any(placeholder in paragraph.text for placeholder in placeholders):
                return child
        return None

    def _paragraphs_before_content_anchor(self, document: DocxDocument, bookmark_names: list[str]) -> list[Paragraph]:
        paragraphs: list[Paragraph] = []
        bookmark_set = set(bookmark_names)
        for paragraph in document.paragraphs:
            if any(start.get(qn("w:name")) in bookmark_set for start in paragraph._p.xpath(".//w:bookmarkStart")):
                break
            paragraphs.append(paragraph)
        return paragraphs

    def _replace_in_paragraph(self, paragraph: Paragraph, replacements: dict[str, str]) -> int:
        if not replacements or not paragraph.text:
            return 0
        count = 0
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
        remaining = paragraph.text
        if not any(old in remaining for old in replacements):
            return count

        new_text = remaining
        for old, new in replacements.items():
            if old in new_text:
                new_text = new_text.replace(old, new)
                count += 1
        self._set_paragraph_plain_text(paragraph, new_text)
        return count

    def _replace_regex_in_paragraph(self, paragraph: Paragraph, pattern: str, replacement: str) -> int:
        if not paragraph.text:
            return 0
        new_text, count = re.subn(pattern, replacement, paragraph.text)
        if count:
            self._set_paragraph_plain_text(paragraph, new_text)
        return count

    def _set_paragraph_plain_text(self, paragraph: Paragraph, text: str) -> None:
        if not paragraph.runs:
            paragraph.add_run(text)
            return
        for run in paragraph.runs:
            run.text = ""
        paragraph.runs[0].text = text

    def _set_bookmark_text_same_paragraph(self, paragraph: Paragraph, start, value: str) -> bool:
        bookmark_id = start.get(qn("w:id"))
        parent = start.getparent()
        if parent is not paragraph._p:
            return False
        children = list(paragraph._p)
        try:
            start_index = children.index(start)
        except ValueError:
            return False
        end_index = None
        for index in range(start_index + 1, len(children)):
            child = children[index]
            if child.tag == qn("w:bookmarkEnd") and child.get(qn("w:id")) == bookmark_id:
                end_index = index
                break
        if end_index is None:
            return False
        for child in children[start_index + 1 : end_index]:
            if child.tag == qn("w:r"):
                for text_node in child.xpath(".//w:t"):
                    text_node.text = ""
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = value
        run.append(text)
        paragraph._p.insert(start_index + 1, run)
        return True

    def _iter_all_paragraphs(self, document: DocxDocument) -> Iterable[Paragraph]:
        for paragraph in document.paragraphs:
            yield paragraph
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
        for section in document.sections:
            for story in (section.header, section.footer, section.first_page_header, section.first_page_footer):
                for paragraph in story.paragraphs:
                    yield paragraph
                for table in story.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            yield from cell.paragraphs

    @staticmethod
    def _format_property(template: str, metadata: dict[str, str]) -> str:
        if template in metadata:
            return metadata.get(template, "")
        try:
            return template.format(**metadata)
        except KeyError:
            return template

    @staticmethod
    def _safe_format(template: str, metadata: dict[str, str]) -> str:
        class SafeDict(dict):
            def __missing__(self, key):
                return ""

        return template.format_map(SafeDict(metadata))
